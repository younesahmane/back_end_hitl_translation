import os
from datetime import datetime
from docx import Document

def export_docx(doc_id: str, segments: list[dict], exports_dir: str) -> str:
    os.makedirs(exports_dir, exist_ok=True)

    doc = Document()
    doc.add_heading(f"Translated Document ({doc_id})", level=1)
    doc.add_paragraph(f"Exported at: {datetime.utcnow().isoformat()}Z")

    doc.add_paragraph("")  # spacer

    for seg in sorted(segments, key=lambda s: s.get("index", 0)):
        tgt = seg.get("human_edit", "") or ""
        # If human_edit empty, fallback to model1
        if not tgt.strip():
            tgt = seg.get("llm_outputs", {}).get("model1", "").strip()

        # Add only translated content
        if tgt.strip():
            doc.add_paragraph(tgt)

    filename = f"{doc_id}_translated.docx"
    outpath = os.path.join(exports_dir, filename)
    doc.save(outpath)
    return filename
